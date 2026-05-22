#!/usr/bin/env python3
"""光伏项目评估报告 Word 生成脚本（详细版）
使用方法：修改 REPORT_DATA 字典后直接运行
依赖：python-docx (pip install python-docx --break-system-packages)
Logo：~/.openclaw/workspace/skills/wind-power-analysis/assets/jianeng_logo_header.png

光伏特有参数：
  - 经营期 25 年 / 折旧 20 年（后5年无折旧）
  - 逆变器替换第12年一次性全部替换（0.05元/W）
  - 组件衰减率 0.5%/年
版本匹配：report_template.md（光伏详细结构）
"""
import os, sys, io, tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image
from lxml import etree

# ───────────────────────────────────────────────
# >>> 配置（每个项目修改此处）<<<
# ───────────────────────────────────────────────
LOGO = os.path.expanduser("~/.openclaw/workspace/skills/wind-power-analysis/assets/jianeng_logo_header.png")
REPORT_DATA = {
    # ── 项目基本信息 ──
    "project_name": "德州",
    "province": "山东",
    "city": "德州",
    "capacity_mw": 100,
    "capacity_side": "直流侧",
    "project_type": "集中式地面电站",
    "panel_type": "N型TOPCon",
    "annual_hours": 1300,
    "degradation_rate": 0.005,
    "curtailment_rate": 0.03,
    "curtailment_detail": "消纳友好省，保守取3%",
    "curtailment_source": "全国新能源消纳监测预警中心2025",
    "report_date": "2026-05-13",

    # ── 限电率与消纳 ──
    "pv_utilization": "98.5%",
    "pv_penetration": "45%",
    "midday_peak_assessment": "午间调峰能力充足，省内负荷支撑良好",
    "curtailment_assessment": "消纳能力≥年发电量1.5倍，无断面阻塞",
    "bottleneck": "无严重断面阻塞，项目位于消纳较好区域",

    # ── 机制电价 ──
    "mechanism_price": 0.360,
    "mechanism_ratio": 0.80,
    "mechanism_years": 12,
    "mechanism_source": "广东省2025年首轮光伏竞价",
    "cfd_note": "实际收入=节点现货+CfD差价补偿（136号文）",

    # ── 分段电价数据 ──
    "mechanism_spot_price": 0.300,
    "mechanism_long_price": 0.380,
    "mechanism_mkt_ratio": 0.70,
    "market_spot_price": 0.250,
    "market_long_price": 0.350,
    "market_mkt_ratio": 0.60,
    "penalty": 0.015,
    "market_fee_rate": 0.02,
    "green_premium": 0.030,

    # ── 核心假设 ──
    "operating_years": 25,
    "depreciation_years": 20,
    "residual_rate": 0.05,
    "om_rates": "1-5年 0.02→6-10年 0.025→11-25年 0.03",
    "inverter_replace_year": 12,
    "inverter_replace_cost": 0.05,
    "interest_rate": 0.04,
    "loan_years": 15,
    "vat_rate": 0.13,
    "om_vat_rate": 0.06,
    "income_tax_rate": 0.25,
    "wacc": 0.055,

    # ── 测算结果 ──
    "t1_limit": 3.85,
    "t4_limit": 4.52,
    "t1_irr": 6.50,
    "t4_irr": 6.00,
    "t4_equity_irr": 12.50,
    "t1_dscr": 1.2000,
    "t4_dscr": 1.2500,
    "t1_lcoe": 0.3200,
    "t4_lcoe": 0.3500,
    "t1_avg_ncf": 520,
    "t1_avg_np": 380,
    "t4_avg_ncf": 480,
    "t4_avg_np": 350,
    "effective_price": 0.3067,

    # ── 分段收益分析 ──
    "mechanism_period_years": 12,
    "mechanism_eff_price": 0.3067,
    "market_period_years": 13,
    "market_eff_price": 0.2695,
    "t1_mech_avg_np": 420,
    "t1_mkt_avg_np": 340,
    "t4_mech_avg_np": 390,
    "t4_mkt_avg_np": 310,

    # ── 决策建议 ──
    "rec_buy_limit": 3.85,
    "rec_sell_limit": 4.52,
    "risk_notes": "光伏午间限电恶化风险（渗透率攀升）、机制电价到期后市场化电价下行风险、逆变器替换第12年现金流凹陷",

    # ── 财务附表数据（关键年份） ──
    "profit_table": [
        [1, 4881.33, 200.00, 61.98, 1545.37, 1301.36, 15.80, 439.21, 1317.62],
        [5, 4784.43, 200.00, 49.61, 1545.37, 954.33, 19.80, 503.83, 1511.50],
        [10, 4666.01, 250.00, 34.16, 1545.37, 520.54, 24.47, 572.87, 1718.60],
        [12, 4619.47, 800.00, 27.98, 1545.37, 347.03, 22.54, 469.14, 1407.42],
        [15, 4550.52, 300.00, 18.71, 1545.37, 86.76, 29.19, 642.63, 1927.88],
        [20, 4437.89, 300.00, 3.25, 1545.37, 0.00, 62.73, 631.64, 1894.91],
        [25, 4326.56, 300.00, 3.25, 1545.37, 0.00, 61.12, 604.19, 1812.56],
    ],
    "fcf_table": [
        [0, "-35,200.00"], [1, "3,186.57"], [5, "3,196.78"],
        [10, "3,032.36"], [15, "2,956.54"], [20, "2,951.05"], [25, "2,940.00"],
    ],
    "equity_cf_table": [
        [0, "—", "—", "—", "-7,040.00"],
        [1, 1317.62, 1545.37, 1560.45, 1302.54],
        [5, 1511.50, 1545.37, 1560.45, 1496.42],
        [10, 1718.60, 1545.37, 1560.45, 1703.52],
        [15, 1927.88, 1545.37, 1560.45, 1912.80],
        [20, 1894.91, 1545.37, 0.00, 3440.28],
        [25, 1812.56, 1545.37, 0.00, 3357.93],
    ],
    "dscr_table": [
        [1, 3721.79, 10.58, 262.35, 3448.86, 1560.45, 347.03, 1907.48, 1.81],
        [5, 3733.46, 14.75, 331.07, 3387.65, 1560.45, 254.49, 1814.94, 1.87],
        [10, 3508.05, 16.22, 357.89, 3133.94, 1560.45, 138.81, 1699.26, 1.84],
        [15, 3402.64, 19.56, 414.25, 2968.82, 1560.45, 23.14, 1583.59, 1.87],
        [16, 3411.39, 62.73, 631.64, 2717.01, 0.00, 0.00, 0.00, float('inf')],
        ["17~25", "3,414~4,260", "61~64", "605~627", "2,717~3,569", 0, 0, 0, "∞"],
    ],
}

TODAY = "20260513"
# ───────────────────────────────────────────────

COLOR_DB = RGBColor(0x1B, 0x3A, 0x5C)

def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = COLOR_DB; r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.font.size = Pt(15)

def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = COLOR_DB; r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.font.size = Pt(13)

def h3(doc, text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.color.rgb = COLOR_DB; r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.font.size = Pt(11)

def para(doc, text, bold=False, sz=11, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.name = '仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    r.font.size = Pt(sz); r.bold = bold

def code_para(doc, text, sz=9):
    """等宽字体段落，用于展示计算过程"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(sz)

def table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C" w:val="clear"/>')
        c._element.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9); r.font.name = '仿宋'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

def add_header(doc, logo_path):
    """江能 Logo 页眉，2cm 宽居右"""
    for s in doc.sections:
        hdr = s.header
        hdr.is_linked_to_previous = False
        p = hdr.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(2))

def add_watermark(doc, logo_path, opacity=0.12):
    """在文档所有页面添加半透明Logo水印（防盗版）"""
    img = Image.open(logo_path)
    rgba = img.convert('RGBA')
    rgba = rgba.resize((360, 300), Image.LANCZOS)
    r, g, b, a = rgba.split()
    a = a.point(lambda x: min(x, int(255 * opacity)))
    watermark = Image.merge('RGBA', (r, g, b, a))
    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    watermark.save(tmp.name, 'PNG')
    tmp_path = tmp.name

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(tmp_path, width=Cm(7))

        for drawing in run._element.findall('.//' + qn('w:drawing')):
            inline = drawing.find(qn('wp:inline'))
            if inline is not None:
                extent = inline.find(qn('wp:extent'))
                cx = extent.get('cx', '0') if extent is not None else '0'
                cy = extent.get('cy', '0') if extent is not None else '0'

                anchor = etree.SubElement(drawing, qn('wp:anchor'))
                anchor.set('simplePos', '0')
                anchor.set('relativeHeight', '0')
                anchor.set('behindDoc', '1')
                anchor.set('locked', '0')
                anchor.set('layoutInCell', '1')
                anchor.set('allowOverlap', '1')

                sp = etree.SubElement(anchor, qn('wp:simplePos'))
                sp.set('x', '0'); sp.set('y', '0')

                hp = etree.SubElement(anchor, qn('wp:positionH'))
                hp.set('relativeFrom', 'page')
                etree.SubElement(hp, qn('wp:align')).text = 'center'

                vp = etree.SubElement(anchor, qn('wp:positionV'))
                vp.set('relativeFrom', 'page')
                etree.SubElement(vp, qn('wp:align')).text = 'center'

                a_ext = etree.SubElement(anchor, qn('wp:extent'))
                a_ext.set('cx', cx); a_ext.set('cy', cy)

                ee = etree.SubElement(anchor, qn('wp:effectExtent'))
                ee.set('l', '0'); ee.set('t', '0'); ee.set('r', '0'); ee.set('b', '0')

                etree.SubElement(anchor, qn('wp:wrapNone'))

                for child in list(inline):
                    tag_local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag_local in ('graphic', 'extent'):
                        anchor.append(child)

                drawing.remove(inline)

    try: os.unlink(tmp_path)
    except: pass


def add_disclaimer(doc):
    """免责声明 — 必须包含"""
    h1(doc, "免责声明")
    text = (
        "本报告由江能能源开发的AI辅助评估系统生成，并经过人工校核，不构成投资建议。"
        "报告中所引用的电力市场数据（机制电价、中长期交易电价、现货电价、限电率、利用小时数等）均来自公开渠道，"
        "部分电价参数在缺乏直接交易数据的情况下采用保守假设推算，可能与实际成交价格存在偏差。"
        "实际项目投资决策应结合以下因素综合判断：\n\n"
        "① 场址实测光照资源数据（至少一个完整年度的Solargis/NASA反演数据）；"
        "② 电网接入批复及送出工程条件；"
        "③ EPC 招标实际报价；"
        "④ 项目所在地最新的机制电价竞价结果；"
        "⑤ 所在省电力交易中心公布的全年现货与中长期交易结算数据。\n\n"
        "本报告中的财务模型基于特定假设（融资利率 4%、经营期 25 年、折旧 20 年、等额本金还款等），"
        "不同融资结构、利率环境及政策变化可能导致测算结果显著偏离。"
        "报告中的'投资边界'和'出售边界'为理论测算阈值，不代表项目实际可实现的交易价格或融资条件。\n\n"
        "江能能源及报告编制方不对因使用本报告而产生的任何直接或间接损失承担责任。"
        "未经授权，不得转载或用于商业用途。"
    )
    para(doc, text, sz=9)

def build(doc, R):
    """构建详细报告正文"""
    gen_mwh = R['capacity_mw'] * R['annual_hours'] * (1 - R['curtailment_rate'])

    # ── 机制期有效电价计算 ──
    mech_wavg = R['mechanism_ratio'] * R['mechanism_price'] + (1 - R['mechanism_ratio']) * R['mechanism_spot_price']
    mech_eff = (mech_wavg + R['green_premium'] - R['penalty']) * (1 - R['market_fee_rate'])
    # ── 市场化期有效电价计算 ──
    mkt_wavg = R['market_mkt_ratio'] * R['market_long_price'] + (1 - R['market_mkt_ratio']) * R['market_spot_price']
    mkt_eff = (mkt_wavg + R['green_premium'] - R['penalty']) * (1 - R['market_fee_rate'])

    # ═══════════ 标题 ═══════════
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"江能能源 · {R['project_name']}{R['capacity_mw']}MW光伏项目投资评估报告")
    r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = COLOR_DB
    r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    para(doc,
         f"评估日期：{R['report_date']} | 分析师：江能研究院（Claude Opus辅助）"
         f"\n项目类型：{R['project_type']} | 装机容量：{R['capacity_mw']}MW（{R['capacity_side']}）"
         f"\n组件类型：{R['panel_type']} | 等效利用小时数：{R['annual_hours']}h（Solargis最佳倾角减100h）"
         f"\n组件衰减率：{R['degradation_rate']*100:.1f}%/年",
         sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ═══════════ 一、电力市场数据 ═══════════
    h1(doc, "一、项目所在地电力市场数据")

    h2(doc, "限电率与消纳情况")
    table(doc, ["指标", "数值", "来源"], [
        ["当地光伏利用率", R['pv_utilization'], R['curtailment_source']],
        ["光伏渗透率", R['pv_penetration'], R['curtailment_source']],
        ["午间调峰能力评估", R['midday_peak_assessment'], "行业分析"],
        ["消纳能力评估", R['curtailment_assessment'], "综合判断"],
        ["断面阻塞情况", R['bottleneck'], "综合判断"],
        ["本项目采用限电率", f"{R['curtailment_rate']*100:.0f}%", "轨A×轨B×隐含限电率取最高"],
    ])

    h2(doc, "机制电价竞价信息（如适用）")
    table(doc, ["项目", "数值", "来源"], [
        ["机制电价", f"{R['mechanism_price']} 元/kWh", R['mechanism_source']],
        ["机制电量比例上限", f"{R['mechanism_ratio']*100:.0f}%", "省级实施细则"],
        ["执行期限", f"{R['mechanism_years']}年", "省级实施细则"],
        ["CfD差价机制", R['cfd_note'], "136号文"],
    ])

    h2(doc, "电价数据（分段）")

    h3(doc, "机制期内（前N年）")
    table(doc, ["电量类型", "权重", "电价", "说明"], [
        ["机制电量", f"{R['mechanism_ratio']*100:.0f}%", f"{R['mechanism_price']:.3f} 元/kWh", "竞价中标价"],
        ["非机制电量（现货）", f"{(1-R['mechanism_ratio'])*100:.0f}%", f"{R['mechanism_spot_price']:.3f} 元/kWh", "现货市场"],
        ["两细则考核", "—", "-0.015 元/kWh", "行业通用"],
        ["市场费用分摊", "—", f"结算电费×{R['market_fee_rate']*100:.0f}%", "—"],
        ["**机制期有效电价**", "—", f"**{mech_eff:.4f} 元/kWh**", "加权后扣除"],
    ])

    h3(doc, "机制到期后（后M年）")
    table(doc, ["电量类型", "权重", "电价", "说明"], [
        ["中长期交易", f"{R['market_mkt_ratio']*100:.0f}%", f"{R['market_long_price']:.3f} 元/kWh", "年度长协"],
        ["现货市场", f"{(1-R['market_mkt_ratio'])*100:.0f}%", f"{R['market_spot_price']:.3f} 元/kWh", "现货市场"],
        ["两细则考核", "—", "-0.015 元/kWh", "行业通用"],
        ["市场费用分摊", "—", f"结算电费×{R['market_fee_rate']*100:.0f}%", "—"],
        ["**市场化期有效电价**", "—", f"**{mkt_eff:.4f} 元/kWh**", "加权后扣除"],
    ])

    para(doc, "⚠️ 机制电价是CfD差价合约，不是固定保底价。实际收入=节点现货+ CfD差价补偿。", bold=True, sz=10)

    # ═══════════ 二、财务测算核心假设 ═══════════
    h1(doc, "二、财务测算核心假设")
    table(doc, ["参数", "数值"], [
        ["装机容量", f"{R['capacity_mw']} MW"],
        ["等效利用小时数", f"{R['annual_hours']}h（Solargis最佳倾角减100h）"],
        ["限电率", f"{R['curtailment_rate']*100:.0f}%"],
        ["组件年衰减率", f"{R['degradation_rate']*100:.1f}%/年"],
        ["首年净发电量", f"{gen_mwh:,.0f} MWh"],
        ["经营期", f"**{R['operating_years']}年**"],
        ["折旧", f"**{R['depreciation_years']}年**直线法，残值率{R['residual_rate']*100:.0f}%（后5年无折旧）"],
        ["运维费", R['om_rates'] + " 元/W·年"],
        ["逆变器替换", f"第{R['inverter_replace_year']}年{R['inverter_replace_cost']}元/W（一次性全部替换，仅此一次）"],
        ["融资利率", f"{R['interest_rate']*100:.0f}%"],
        ["融资期限", f"**{R['loan_years']}年**"],
        ["还款方式", "等额本金"],
        ["增值税", f"{R['vat_rate']*100:.0f}%（销项），进项含融资租赁租金×13%+运维费×6%"],
        ["企业所得税", f"{R['income_tax_rate']*100:.0f}%"],
        ["全投资WACC", f"{R['wacc']*100:.1f}%"],
    ])

    # ═══════════ 三、财务测算结果 ═══════════
    h1(doc, "三、财务测算结果")

    h2(doc, "3.1 模型边界条件")
    table(doc, ["参数", "数值"], [
        ["装机容量", f"{R['capacity_mw']} MW"],
        ["等效利用小时数", f"{R['annual_hours']} h"],
        ["限电率", f"{R['curtailment_rate']*100:.1f}%"],
        ["年净发电量", f"{gen_mwh:,.0f} MWh"],
        ["有效电价", f"{R['effective_price']} 元/kWh"],
        ["经营期", f"{R['operating_years']} 年"],
        ["融资利率", f"{R['interest_rate']*100:.0f}%"],
        ["融资期限", f"{R['loan_years']} 年"],
        ["折旧年限", f"{R['depreciation_years']} 年"],
        ["残值率", f"{R['residual_rate']*100:.0f}%"],
        ["运维费第1段(1-5年)", f"0.02 元/W·年"],
        ["运维费第2段(6-10年)", f"0.025 元/W·年"],
        ["运维费第3段(11-25年)", f"0.03 元/W·年"],
        ["逆变器替换(第12年)", f"0.05 元/W"],
        ["增值税率", f"{R['vat_rate']*100:.0f}%"],
        ["所得税率", f"{R['income_tax_rate']*100:.0f}%"],
    ])

    h2(doc, "3.2 任务1：100%融资，最小DSCR≥1.2")
    table(doc, ["指标", "数值"], [
        ["最高单瓦投资（元/W）", f"{R['t1_limit']}"],
        ["总投资（亿元）", f"{R['t1_limit']*R['capacity_mw']/100:.2f}"],
        ["度电成本（元/kWh）", f"{R['t1_lcoe']:.4f}"],
        ["全投资IRR（%）", f"{R['t1_irr']:.2f}"],
        ["最小DSCR", f"{R['t1_dscr']:.4f}"],
        [f"{R['operating_years']}年平均净现金流量（万元）", f"{R['t1_avg_ncf']:.2f}"],
        [f"{R['operating_years']}年平均净利润（万元）", f"{R['t1_avg_np']:.2f}"],
    ])

    h2(doc, "3.3 任务4：80%融资，全投资IRR≥6%且资本金IRR≥8%")
    table(doc, ["指标", "数值"], [
        ["目标单瓦投资（元/W）", f"{R['t4_limit']}"],
        ["总投资（亿元）", f"{R['t4_limit']*R['capacity_mw']/100:.2f}"],
        ["度电成本（元/kWh）", f"{R['t4_lcoe']:.4f}"],
        ["全投资IRR（%）", f"{R['t4_irr']:.2f}"],
        ["税后资本金IRR（%）", f"{R['t4_equity_irr']:.2f}"],
        ["最小DSCR", f"{R['t4_dscr']:.4f}"],
        [f"{R['operating_years']}年平均净现金流量（万元）", f"{R['t4_avg_ncf']:.2f}"],
        [f"{R['operating_years']}年平均净利润（万元）", f"{R['t4_avg_np']:.2f}"],
    ])

    # ═══════════ 四、财务指标附表 ═══════════
    h1(doc, "四、财务指标附表")

    h2(doc, "表1：资本金投资利润表（出售边界·80%融资）（单位：万元）")
    p1_h = ["年份","营业收入","运维费","保险费","折旧","利息","增值税及附加","所得税","净利润"]
    p1_r = [[str(r[0]), f"{r[1]:.2f}", f"{r[2]:.2f}", f"{r[3]:.2f}", f"{r[4]:.2f}", f"{r[5]:.2f}", f"{r[6]:.2f}", f"{r[7]:.2f}", f"{r[8]:.2f}"] for r in R['profit_table']]
    table(doc, p1_h, p1_r)
    para(doc, "说明：营业收入=年发电量×有效电价 | 运维费=装机×单位费率(分段)+第12年逆变器替换 | 保险费=净值×0.2% | 折旧=(总投资×(1-残值率))/20年 | 利息=期初余额×利率 | 增值税=销项税-进项税 | 所得税=(收入-运维-保险-折旧-利息-增值税及附加)×25% | 净利润=营业收入-运维费-保险费-折旧-利息-增值税及附加-所得税", sz=9)

    h2(doc, "表2：全投资净现金流表（出售边界口径）（单位：万元）")
    fcf_h = ["年份", "全投资FCF"]
    fcf_r = [[str(r[0]), str(r[1])] for r in R['fcf_table']]
    table(doc, fcf_h, fcf_r)
    para(doc, f"全投资FCF基于无杠杆附加税（剔除利息税盾）。t=0为初始投资流出{R['t4_limit']*R['capacity_mw']/100:.2f}亿元。全投资IRR={R['t4_irr']:.2f}%。", sz=9)

    h2(doc, "表3：资本金投资现金流表（出售边界·80%融资）（单位：万元）")
    eq_h = ["年份", "净利润", "折旧", "偿还本金", "股权现金流"]
    def fmt_val(v):
        if isinstance(v, str): return v
        return f"{v:,.2f}"
    eq_r = [[str(r[0]), fmt_val(r[1]), fmt_val(r[2]), fmt_val(r[3]), fmt_val(r[4])] for r in R['equity_cf_table']]
    table(doc, eq_h, eq_r)
    para(doc, f"t=0初始资本金投入{R['t4_limit']*R['capacity_mw']/5:,.0f}万元（总投资{R['t4_limit']*R['capacity_mw']/100:.2f}亿×20%）。股权现金流=净利润+折旧-偿还本金。税后资本金IRR={R['t4_equity_irr']:.2f}%。", sz=9)

    h2(doc, "表4：偿债覆盖计算表（投资边界·100%融资）（单位：万元）")
    dscr_h = ["年份","EBITDA","增值税及附加","所得税","可用于还款","应还本金","应还利息","应还本息","DSCR"]
    dscr_r = []
    for r in R['dscr_table']:
        row = [str(r[0])]
        for v in r[1:]:
            if isinstance(v, str):
                row.append(v)
            elif isinstance(v, float) and (v == float('inf') or v == float('inf')):
                row.append("∞")
            else:
                row.append(f"{v:,.2f}" if abs(v) >= 100 else f"{v:.2f}")
        dscr_r.append(row)
    table(doc, dscr_h, dscr_r)
    para(doc, f"说明：DSCR=可用于还款/应还本息。可用于还款=EBITDA-增值税及附加-所得税。应还本金=期初本金/{R['loan_years']}年（等额本金）。最小DSCR={R['t1_dscr']:.2f}，全期偿债能力达标。", sz=9)

    # ═══════════ 五、投资决策建议 ═══════════
    h1(doc, "五、投资决策建议")
    para(doc, f"1. 推荐单瓦投资阈值：投资边界{R['rec_buy_limit']}元/W（100%融资，DSCR≥1.2）；"
              f"出售边界{R['rec_sell_limit']}元/W（80%融资，全投资IRR≥6%且资本金IRR≥8%）", sz=10)
    para(doc, "2. 融资视角：净建造成本若高于投资边界，建议控制成本或引入资本金。", sz=10)
    para(doc, "3. 出售视角：EPC总价若低于出售边界，80%融资下具备出售价值。", sz=10)
    para(doc, f"4. 风险提示：{R['risk_notes']}", sz=10)

    # ═══════════ 免责声明 ═══════════
    add_disclaimer(doc)


def generate(R):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
        s.header_distance = Cm(1.0)
    add_header(doc, LOGO)
    build(doc, R)
    add_watermark(doc, LOGO)
    out = os.path.expanduser(
        f"~/.openclaw/workspace/projects/{R['project_name']}{R['capacity_mw']}MW光伏/"
        f"{R['project_name']}{R['capacity_mw']}MW光伏_评估报告_{TODAY}.docx"
    )
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    print(f"✓ {out}")

if __name__ == '__main__':
    generate(REPORT_DATA)
